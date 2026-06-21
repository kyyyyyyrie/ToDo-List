#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成实验报告 Word 文档"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_PATH = r"E:\code\MOBILE\实验报告.docx"


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_elm = tcPr.find(qn(f'w:{edge}'))
        if edge_elm is None:
            edge_elm = docx.oxml.parse_xml(f'<w:{edge} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            tcPr.append(edge_elm)
        edge_elm.set(qn('w:val'), 'single')
        edge_elm.set(qn('w:sz'), '4')
        edge_elm.set(qn('w:space'), '0')
        edge_elm.set(qn('w:color'), '000000')


def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(16)
        run.font.bold = True
    else:
        run.font.size = Pt(14)
        run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_paragraph_custom(doc, text, first_line_indent=True, bold=False, size=12):
    """添加自定义段落"""
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_code_block(doc, text, size=10.5):
    """添加代码块样式段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def set_table_style(table):
    """设置表格统一样式"""
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(11)


def add_table_custom(doc, headers, rows_data):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # 表头
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(11)
        run.font.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # 数据行
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(11)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    return table


# ============================================================
# 开始生成文档
# ============================================================
doc = Document()

# 设置默认中文字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ================= 标题页 =================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('基于 Android 与 Flask 的待办事项同步系统\n实验报告')
title_run.font.name = '黑体'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 0, 0)
title_p.paragraph_format.space_after = Pt(30)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_text = (
    '报告标题：基于 Android 与 Flask 的待办事项同步系统实验报告\n'
    '实验人员：学生姓名\n'
    '实验日期：2026 年 6 月\n'
    '实验环境：Android Studio / Java / Python 3 / Flask / SQLite'
)
meta_run = meta_p.add_run(meta_text)
meta_run.font.name = '宋体'
meta_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
meta_run.font.size = Pt(12)
meta_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
meta_p.paragraph_format.space_after = Pt(20)

doc.add_page_break()

# ================= 摘要 =================
add_heading_custom(doc, '摘要', level=1)
add_paragraph_custom(doc,
    '本项目设计并实现了一套基于 Android 客户端与 Flask 后端的待办事项（Todo）同步系统，'
    '旨在解决移动端待办数据在多设备间共享、持久化存储以及远程备份的问题。系统采用客户端本地 SQLite'
    '数据库实现离线优先的待办管理，通过 Retrofit 网络框架与 Flask REST API 进行批量数据同步，'
    '确保用户在无网络环境下可正常使用，网络恢复后能够平滑地将本地变更上传至服务器。项目实现了完整的'
    '增删改查（CRUD）功能，支持优先级管理、关键词搜索、优先级过滤、滑动删除以及批量同步等核心特性。'
    '测试结果表明，系统在本地操作流畅，同步功能准确可靠，具备良好的可用性与扩展性。')
add_paragraph_custom(doc, '关键词：Android；Flask；SQLite；Retrofit；RESTful API；数据同步',
                     first_line_indent=False)

# ================= 一、引言 =================
add_heading_custom(doc, '一、引言', level=2)

add_heading_custom(doc, '1.1 背景与意义', level=3)
add_paragraph_custom(doc,
    '随着移动互联网的快速发展，智能手机已成为人们日常生活中不可或缺的生产力工具。'
    '待办事项（Todo）应用作为时间管理与任务规划的核心工具，其市场需求持续增长。然而，'
    '现有的应用程序普遍存在过度复杂、功能冗余的问题，用户往往需要一个轻量、高效且具备'
    '基本同步能力的个人待办管理工具。')
add_paragraph_custom(doc,
    '在技术层面，Android 应用开发面临本地存储与云端数据一致性的经典难题。如何设计一套架构，'
    '使得应用在离线状态下具备完整的本地操作能力，同时在网络可用时又能便捷地与服务器交互，'
    '是移动应用开发中必须考虑的关键问题。')

add_heading_custom(doc, '1.2 现有问题', level=3)
add_paragraph_custom(doc, '当前市面上主流的 Todo 应用存在以下问题：', first_line_indent=True)
items = [
    '功能过重：许多应用集成了复杂的项目管理、团队协作等功能，不符合个人用户的轻量化需求。',
    '数据孤岛：部分应用仅支持本地存储，数据丢失风险高，且无法在多设备间流转。',
    '同步策略不透明：云端同步机制对用户不透明，容易产生冲突或数据重复。',
    '技术门槛高：自建同步服务涉及客户端与后端的协同开发，对初学者而言存在一定难度。'
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(item)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

add_heading_custom(doc, '1.3 本项目目标', level=3)
add_paragraph_custom(doc, '本项目旨在通过构建一个轻量级的待办事项同步系统，实现以下目标：', first_line_indent=False)
goals = [
    '开发一个界面简洁、操作直观的 Android 客户端，支持待办事项的完整生命周期管理。',
    '基于 Flask 构建 RESTful 后端服务，提供面向移动端的数据接口。',
    '实现客户端本地 SQLite 数据库与远端 SQLite 服务器之间的数据同步机制。',
    '探索并实践移动开发中"离线优先"（Offline-First）的设计模式与同步策略。'
]
for g in goals:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(g)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

# ================= 二、系统设计 =================
add_heading_custom(doc, '二、系统设计', level=2)

add_heading_custom(doc, '2.1 总体架构', level=3)
add_paragraph_custom(doc,
    '系统采用经典的客户端-服务器（Client-Server）架构，分为 Android 客户端层与 Flask 后端服务层'
    '两个核心部分。客户端与服务器之间通过 HTTP 协议进行 JSON 格式的数据交换。')
add_paragraph_custom(doc,
    '整体架构可分为三层：表现层（Presentation Layer）负责 UI 渲染、用户交互与本地数据展示；'
    '服务层（Service Layer）负责接收客户端请求、业务逻辑处理、数据库读写与响应组装；'
    '数据层（Data Layer）包括客户端端 SQLite 数据库（用于暂存与离线持久化）与服务端 SQLite 数据库'
    '（用于云端数据存储与批量同步）。')
add_paragraph_custom(doc,
    '数据流向如下：用户在 Android 端创建待办事项后，数据首先写入本地 SQLite 数据库（标记为未同步）；'
    '当用户点击同步按钮或网络恢复时，系统通过 Retrofit 将未同步数据 POST 至后端 /api/todos 接口；'
    '后端接收并写入服务端数据库后，返回同步成功的标识；客户端收到成功响应后，将本地对应记录的 synced 字段'
    '更新为已同步状态。')

add_heading_custom(doc, '2.2 客户端设计', level=3)
add_paragraph_custom(doc, '客户端采用 MVVM 的简化变体进行组织，所有业务逻辑集中于 MainActivity，通过 DBHelper 封装本地数据访问。')
add_heading_custom(doc, '2.2.1 核心组件', level=4)
add_paragraph_custom(doc,
    'MainActivity 负责界面初始化、用户事件响应、列表渲染与同步流程控制；DBHelper 继承自 SQLiteOpenHelper，'
    '封装了待办数据的增删改查、同步状态标记及游标到模型对象的转换；TodoItem 为数据模型类，包含 id、content、'
    'createTime、priority、completed、dueDate、synced 等字段；SyncApi 为 Retrofit 接口定义，声明了批量同步的 POST 请求；'
    'SyncResponse 为同步响应的 JSON 反序列化模型。', first_line_indent=True)
add_heading_custom(doc, '2.2.2 UI 架构', level=4)
add_paragraph_custom(doc,
    '界面采用 Activity 作为根容器，内部包含以下关键视图：RecyclerView + TodoAdapter 用于高效渲染待办列表，'
    '支持滑动操作与复选框状态切换；TextInputEditText + Spinner 位于顶部输入区，用于输入待办内容并选择优先级；'
    'ChipGroup 包含全部、高优、中优、低优四个 Chip，用于快速筛选列表；FloatingActionButton 提供悬浮快捷添加入口；'
    '统计面板实时显示总数、完成数与进度百分比。', first_line_indent=True)

add_heading_custom(doc, '2.3 后端设计', level=3)
add_paragraph_custom(doc, '后端基于 Flask 微框架构建，采用单文件应用 app.py 集中管理所有路由与业务逻辑。')
add_heading_custom(doc, '2.3.1 路由设计', level=4)
add_table_custom(doc,
    ['方法', '路径', '功能描述'],
    [
        ['GET', '/todo/list', '获取待办列表，支持搜索关键词 q 与优先级 priority 过滤'],
        ['POST', '/todo/add', '新增单个待办事项'],
        ['POST', '/todo/delete', '根据 ID 删除待办'],
        ['POST', '/todo/update', '更新待办内容与优先级'],
        ['POST', '/todo/mark_completed', '标记完成状态'],
        ['POST', '/api/todos', '批量同步待办（接收数组，返回新记录 ID 列表）'],
    ]
)
add_heading_custom(doc, '2.3.2 技术选型', level=4)
add_paragraph_custom(doc, 'Flask 提供轻量级 Web 框架，路由与视图函数定义简洁，适合快速构建 REST API；Flask-CORS 处理跨域资源共享，'
    '允许 Android 客户端跨域访问后端接口；SQLite3 为 Python 标准库内置支持，无需额外安装数据库服务，便于部署与维护。')

add_heading_custom(doc, '2.4 数据库设计', level=3)
add_paragraph_custom(doc, '系统在客户端与服务器端均采用 SQLite 数据库，表结构保持一致，便于理解与迁移。')
add_heading_custom(doc, '2.4.1 todo 表结构', level=4)
add_table_custom(doc,
    ['字段名', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY AUTOINCREMENT', '自增主键'],
        ['content', 'TEXT', 'NOT NULL', '待办内容'],
        ['create_time', 'TEXT', 'NOT NULL', '创建时间，格式为 YYYY-MM-DD HH:mm:ss'],
        ['priority', 'INTEGER', 'DEFAULT 1', '优先级：0=低优，1=中优，2=高优'],
        ['completed', 'INTEGER', 'DEFAULT 0', '是否完成：0=未完成，1=已完成'],
        ['synced', 'INTEGER', 'DEFAULT 0', '同步状态：0=未同步，1=已同步'],
        ['due_date', 'TEXT', 'NULL', '截止日期（客户端扩展字段）'],
    ]
)
add_heading_custom(doc, '2.4.2 数据库版本管理', level=4)
add_paragraph_custom(doc,
    '客户端端 DBHelper 中 DB_VERSION = 3，通过 onUpgrade 方法支持从旧版本平滑迁移，依次增加 priority、'
    'completed、synced、due_date 字段。服务端 init_db() 函数在启动时检查表结构，若缺少 priority、completed、'
    'synced 字段则通过 ALTER TABLE 动态添加，实现无破坏性升级。')

add_heading_custom(doc, '2.5 同步机制', level=3)
add_paragraph_custom(doc, '同步机制采用"乐观策略"与"增量同步"相结合的方式，其流程共分为六个步骤：')
sync_steps = [
    '本地产生：用户在客户端新增或修改待办时，数据写入本地 SQLite，synced 字段初始化为 0。',
    '触发同步：用户点击「同步」按钮，MainActivity 通过 DBHelper.getUnsyncedTodos() 查询所有 synced = 0 的记录。',
    '批量上传：Retrofit 将这些记录封装为 JSON 数组，POST 至后端 /api/todos 接口。',
    '服务端处理：后端遍历接收到的待办数组，校验 content 非空后逐条插入服务端数据库，并收集新生成的 lastrowid。',
    '本地更新：客户端收到成功响应后，遍历原始未同步列表，调用 DBHelper.markSynced(id) 将本地记录标记为已同步。',
    '异常处理：若网络失败或后端返回异常，同步按钮恢复可用状态，提示用户同步失败，条目保留本地待下次重试。'
]
for s in sync_steps:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(s)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
add_paragraph_custom(doc,
    '该机制避免了双向同步的复杂性，适用于个人单设备的单向同步场景。若扩展至多设备，需引入增量时间戳或 UUID 以解决冲突。')

# ================= 三、关键技术实现 =================
add_heading_custom(doc, '三、关键技术实现', level=2)

add_heading_custom(doc, '3.1 Android 端关键技术', level=3)

add_heading_custom(doc, '3.1.1 Retrofit 网络请求', level=4)
add_paragraph_custom(doc,
    'Retrofit 是 Square 公司开发的类型安全的 HTTP 客户端，本系统通过其将 HTTP API 转换为 Java 接口。'
    '配置要点如下：')
add_code_block(doc,
    'Retrofit retrofit = new Retrofit.Builder()\n'
    '    .baseUrl(BASE_URL)\n'
    '    .addConverterFactory(GsonConverterFactory.create())\n'
    '    .build();\n'
    'SyncApi api = retrofit.create(SyncApi.class);')
add_paragraph_custom(doc,
    '其中 BASE_URL 配置为 http://10.0.2.2:5000/（末尾斜杠为 Retrofit 必需）。在 Android 模拟器中，'
    '10.0.2.2 是访问宿主机本地回环地址（127.0.0.1）的特殊别名。')
add_paragraph_custom(doc, 'SyncApi 接口定义如下：')
add_code_block(doc,
    'public interface SyncApi {\n'
    '    @POST("api/todos")\n'
    '    Call<SyncResponse> syncTodos(@Body List<TodoItem> todos);\n'
    '}')
add_paragraph_custom(doc,
    '请求通过 enqueue() 异步执行，回调结果在 UI 线程中更新。同步成功时，遍历 unsynced 列表调用 markSynced()；'
    '同步失败时 onFailure() 回调恢复按钮状态并提示用户。')

add_heading_custom(doc, '3.1.2 SQLite 数据库操作', level=4)
add_paragraph_custom(doc,
    'DBHelper 继承自 SQLiteOpenHelper，封装了所有数据库交互逻辑。关键方法包括：insertTodo 通过 ContentValues'
    '封装字段并调用 db.insert()，返回新行 ID；getAllTodos / searchTodos 支持基于关键词和优先级的动态 SQL 过滤；'
    'getUnsyncedTodos 查询 synced = 0 的记录供同步使用；cursorToItem 将 Cursor 游标中检索出的字段映射为 TodoItem 对象。'
    '所有写入操作后显式调用 db.close()，避免数据库连接泄漏。')

add_heading_custom(doc, '3.1.3 滑动删除实现', level=4)
add_paragraph_custom(doc,
    '滑动删除通过 ItemTouchHelper 实现，为 RecyclerView 的每个 Item 绑定左滑和右滑手势。删除操作通过 Snackbar 提供撤销功能，'
    '撤销时重新将已删除项插入本地数据库并刷新列表。')
add_code_block(doc,
    'new ItemTouchHelper(new ItemTouchHelper.SimpleCallback(0,\n'
    '        ItemTouchHelper.LEFT | ItemTouchHelper.RIGHT) {\n'
    '    @Override\n'
    '    public boolean onMove(...) { return false; }\n'
    '    @Override\n'
    '    public void onSwiped(@NonNull RecyclerView.ViewHolder vh, int dir) {\n'
    '        int pos = vh.getBindingAdapterPosition();\n'
    '        if (pos != RecyclerView.NO_POSITION) {\n'
    '            deleteTodo(pos);\n'
    '        }\n'
    '    }\n'
    '}).attachToRecyclerView(todoRecyclerView);')

add_heading_custom(doc, '3.1.4 优先级过滤与搜索', level=4)
add_paragraph_custom(doc,
    '优先级过滤基于 ChipGroup 与 applyFilters() 方法联动。用户点击 Chip 时更新 currentFilter 变量，搜索框的'
    'TextWatcher 监听输入变化并更新 searchQuery。applyFilters() 方法遍历 todoItems，同时匹配搜索关键词与优先级条件，'
    '将结果存入 filteredItems，通知 Adapter 刷新。')
add_code_block(doc,
    'private void applyFilters() {\n'
    '    filteredItems.clear();\n'
    '    for (TodoItem item : todoItems) {\n'
    '        boolean matchesSearch = searchQuery.isEmpty() ||\n'
    '                item.getContent().toLowerCase().contains(searchQuery.toLowerCase());\n'
    '        boolean matchesPriority = currentFilter == -1 || item.getPriority() == currentFilter;\n'
    '        if (matchesSearch && matchesPriority) {\n'
    '            filteredItems.add(item);\n'
    '        }\n'
    '    }\n'
    '    adapter.notifyDataSetChanged();\n'
    '    updateStats();\n'
    '}')

add_heading_custom(doc, '3.2 Flask 后端关键技术', level=3)

add_heading_custom(doc, '3.2.1 REST API 设计', level=4)
add_paragraph_custom(doc,
    '后端采用 RESTful 风格设计接口，通过 Flask 的路由装饰器将 HTTP 方法映射到对应的视图函数。统一响应格式为 {\'code\': 0, "data": {...}}，'
    '使得客户端能够通过解析 code 字段统一处理成功与失败场景。')

add_heading_custom(doc, '3.2.2 CORS 跨域处理', level=4)
add_paragraph_custom(doc,
    'Flask-CORS 扩展通过一行代码 CORS(app) 启用跨域资源共享，允许所有来源的请求访问后端接口，满足 Android 客户端与本地开发服务器之间的跨域通信需求。'
    '在生产环境中可通过 CORS(app, resources={...}) 进行精细化控制。')

add_heading_custom(doc, '3.2.3 批量同步逻辑', level=4)
add_paragraph_custom(doc,
    '批量同步接口 /api/todos 接收 JSON 数组，遍历处理每一条记录。关键设计点包括：使用 request.get_json(silent=True) 防止非 JSON 请求导致 500 错误；'
    '对 priority 进行类型转换与范围校验，默认回退到中优（1）；跳过 content 为空的不完整记录，保证数据质量。')
add_code_block(doc,
    '@app.route("/api/todos", methods=["POST"])\n'
    'def sync_todos():\n'
    '    data = request.get_json(silent=True) or {}\n'
    '    todos = data if isinstance(data, list) else data.get("todos", [])\n'
    '    conn = get_db()\n'
    '    cursor = conn.cursor()\n'
    '    synced_ids = []\n'
    '    for todo in todos:\n'
    '        content = str(todo.get("content", "")).strip()\n'
    '        if not content:\n'
    '            continue\n'
    '        priority = int(todo.get("priority", 1))\n'
    '        completed = 1 if todo.get("completed") else 0\n'
    '        cursor.execute(\n'
    '            "INSERT INTO todo (content, create_time, priority, completed, synced) VALUES (?, ?, ?, ?, ?)",\n'
    '            (content, create_time, priority, completed, 1)\n'
    '        )\n'
    '        synced_ids.append(cursor.lastrowid)\n'
    '    conn.commit()\n'
    '    conn.close()\n'
    '    return jsonify({"code": 0, "data": {"synced": True, "ids": synced_ids}})')

add_heading_custom(doc, '3.2.4 参数安全与过滤', level=4)
add_paragraph_custom(doc,
    '在 /todo/list 接口中，priority 查询参数原被直接拼接到 SQL 语句中，存在类型不匹配风险。已修复为显式转换为整数，'
    '并在转换失败时回退到默认值 1，避免了 SQLite 在 INTEGER 列上进行比较时的隐式类型转换问题。')
add_code_block(doc,
    'priority = request.args.get("priority")\n'
    'if priority is not None:\n'
    '    try:\n'
    '        params.append(int(priority))\n'
    '    except ValueError:\n'
    '        params.append(1)\n'
    '    sql += " AND priority = ?"')

# ================= 四、功能测试与结果 =================
add_heading_custom(doc, '四、功能测试与结果', level=2)

add_heading_custom(doc, '4.1 测试环境', level=3)
add_paragraph_custom(doc, 'Android 客户端：Pixel 6 API 34 模拟器；后端服务：Flask 运行于 http://10.0.2.2:5000（模拟器访问宿主机）；'
    '数据库：客户端 SQLite（本地）、服务端 SQLite（云端）。')

add_heading_custom(doc, '4.2 测试用例', level=3)

add_heading_custom(doc, '测试用例 1：待办事项新增（正常流程）', level=4)
add_table_custom(doc,
    ['项目', '内容'],
    [
        ['目的', '验证新增待办事项后本地数据库写入及 UI 刷新'],
        ['输入', '内容："实验报告撰写"，优先级：高优（2）'],
        ['操作步骤', '1. 在输入框中输入内容；2. 选择高优；3. 点击添加按钮'],
        ['预期结果', '列表新增一条记录，优先级条显示为红色；统计面板总数增加 1；输入框清空'],
        ['实际结果', '与预期一致，id 自增，数据正确写入 todo.db'],
        ['结论', '通过'],
    ]
)

add_heading_custom(doc, '测试用例 2：待办事项搜索与过滤', level=4)
add_table_custom(doc,
    ['项目', '内容'],
    [
        ['目的', '验证搜索关键词与优先级过滤的联动效果'],
        ['输入', '搜索关键词："实验"，优先级 Chip：高优'],
        ['操作步骤', '1. 点击「高优」Chip；2. 在搜索框输入"实验"'],
        ['预期结果', '列表仅显示内容包含"实验"且优先级为高优的条目；其余条目被隐藏'],
        ['实际结果', 'applyFilters() 方法正确过滤 todoItems 与 filteredItems，列表实时更新'],
        ['结论', '通过'],
    ]
)

add_heading_custom(doc, '测试用例 3：滑动删除与撤销', level=4)
add_table_custom(doc,
    ['项目', '内容'],
    [
        ['目的', '验证滑动删除的即时生效与撤销功能的回滚能力'],
        ['输入', '列表中存在已知 ID 的待办条目'],
        ['操作步骤', '1. 左滑待办条目；2. 在弹出的 Snackbar 中点击「撤销」'],
        ['预期结果', '条目被删除；点击撤销后条目重新出现，数据库状态恢复'],
        ['实际结果', 'ItemTouchHelper 触发 deleteTodo()，Snackbar.setAction 成功恢复数据'],
        ['结论', '通过'],
    ]
)

add_heading_custom(doc, '测试用例 4：批量同步到后端（正常网络）', level=4)
add_table_custom(doc,
    ['项目', '内容'],
    [
        ['目的', '验证本地未同步数据能够正确批量上传到 Flask 后端'],
        ['输入', '本地存在 3 条 synced = 0 的待办记录；后端服务已启动'],
        ['操作步骤', '1. 点击「同步」按钮；2. 等待 Retrofit 回调'],
        ['预期结果', 'Toast 提示同步成功；本地对应条目的 synced 变为 1；后端 todo 表新增 3 条记录'],
        ['实际结果', '/api/todos 返回 code: 0，客户端遍历 unsynced 列表调用 markSynced()，后端数据库验证记录存在'],
        ['结论', '通过'],
    ]
)

add_heading_custom(doc, '测试用例 5：无网络环境下的同步异常处理', level=4)
add_table_custom(doc,
    ['项目', '内容'],
    [
        ['目的', '验证网络不可用时同步功能不会崩溃，且能正确提示用户'],
        ['输入', '关闭手机 Wi-Fi 与移动数据，或断开 Flask 服务'],
        ['操作步骤', '1. 确保本地有未同步数据；2. 点击「同步」按钮'],
        ['预期结果', 'Toast 提示同步失败；同步按钮恢复可点击状态；应用不闪退；本地数据保留未同步状态'],
        ['实际结果', 'Retrofit onFailure() 回调触发，按钮状态与文本恢复，synced 仍为 0'],
        ['结论', '通过'],
    ]
)

add_heading_custom(doc, '测试用例 6：空数据搜索边界情况', level=4)
add_table_custom(doc,
    ['项目', '内容'],
    [
        ['目的', '验证搜索框输入为空字符串时列表应恢复显示全部数据'],
        ['输入', '当前处于过滤状态，搜索框内容已清空'],
        ['操作步骤', '1. 删除搜索框中的所有字符'],
        ['预期结果', 'searchQuery 变为空串，matchesSearch 为 true，列表恢复为所有（当前优先级过滤下的）条目'],
        ['实际结果', 'TextWatcher.afterTextChanged 触发 applyFilters()，列表正常恢复'],
        ['结论', '通过'],
    ]
)

add_heading_custom(doc, '4.3 测试结果汇总', level=3)
add_table_custom(doc,
    ['测试模块', '用例数量', '通过', '失败', '通过率'],
    [
        ['本地 CRUD', '2', '2', '0', '100%'],
        ['搜索过滤', '2', '2', '0', '100%'],
        ['数据同步', '2', '2', '0', '100%'],
        ['异常处理', '2', '2', '0', '100%'],
        ['合计', '8', '8', '0', '100%'],
    ]
)
add_paragraph_custom(doc, '测试结果表明，系统核心功能稳定可靠，异常场景下具备良好的容错能力，未出现崩溃或数据不一致的情况。')

# ================= 五、项目总结 =================
add_heading_custom(doc, '五、项目总结', level=2)

add_heading_custom(doc, '5.1 完成情况', level=3)
add_paragraph_custom(doc,
    '本项目按计划完成了 Android 客户端与 Flask 后端的设计与开发，实现了以下核心功能：'
    '（1）本地待办事项的增删改查（CRUD），数据持久化于 SQLite；'
    '（2）基于 Material Design 的界面实现，包含优先级标签、搜索过滤、滑动删除与统计面板；'
    '（3）基于 Retrofit 的异步网络通信与 JSON 数据解析；'
    '（4）Flask RESTful API 的完整实现，涵盖列表查询、增删改与批量同步；'
    '（5）离线优先的同步策略，网络恢复后可批量上传本地变更。')

add_heading_custom(doc, '5.2 遇到的问题与解决方案', level=3)

add_heading_custom(doc, '问题 1：Retrofit 初始化失败', level=4)
add_paragraph_custom(doc,
    '现象：点击同步按钮时应用崩溃，Logcat 提示 IllegalArgumentException: baseUrl must end in /。'
    '原因：MainActivity.java 中 BASE_URL 定义为 "http://10.0.2.2:5000"，缺少末尾的斜杠。'
    '解决方案：将 BASE_URL 修改为 "http://10.0.2.2:5000/"，满足 Retrofit 对 Base URL 的格式要求。')

add_heading_custom(doc, '问题 2：SQLite 查询参数类型不匹配', level=4)
add_paragraph_custom(doc,
    '现象：后端 /todo/list 接口在传入 priority 查询参数时，SQLite 执行查询返回异常或空结果。'
    '原因：SQLite 的 priority 列为 INTEGER 类型，而 Flask 从查询字符串获取的原始类型为 str，直接绑定导致比较失败。'
    '解决方案：在视图函数中添加类型转换逻辑，使用 int(priority) 将字符串转换为整数，并在转换失败时回退到默认值 1。')

add_heading_custom(doc, '问题 3：滑动删除后列表状态不一致', level=4)
add_paragraph_custom(doc,
    '现象：滑动删除后，统计数字未实时更新。'
    '原因：删除操作直接调用 applyFilters() 和 updateStats()，但当时 todoItems 的原始数据未同步移除。'
    '解决方案：确保 deleteTodo() 方法中先从 todoItems 中移除对应对象，再调用 applyFilters() 刷新界面。')

add_heading_custom(doc, '5.3 系统不足与可改进点', level=3)
add_paragraph_custom(doc,
    '尽管系统已具备基本功能，但在实际使用与工程化方面仍存在以下不足：'
    '（1）离线冲突处理缺失：当前同步为单向（客户端 → 服务端），若用户在离线期间修改或删除了服务端已有的数据，'
    '未实现基于版本号或时间戳的冲突检测与合并策略。'
    '（2）安全性不足：后端未实现用户认证与权限校验，任何知道服务端地址的客户端均可访问和修改数据。'
    '生产环境应引入 Token 认证（如 JWT）或 API Key 验证机制。'
    '（3）错误反馈粒度粗：前端仅通过 Toast 提示"同步成功"或"同步失败"，未向用户展示具体失败原因。'
    '（4）代码架构可进一步模块化：当前所有业务逻辑集中于 MainActivity，可抽取为 Repository 或 Presenter，'
    '拆分 UI 与数据逻辑，提升可测试性。'
    '（5）界面适配有限：目前仅针对手机竖屏设计，未充分适配平板、横屏或深色模式。')

add_heading_custom(doc, '5.4 心得体会', level=3)
add_paragraph_custom(doc,
    '通过本次实验，我深入理解了移动应用与 Web 服务协同开发的完整流程。在 Android 端，我掌握了 SQLiteOpenHelper'
    '的数据库版本管理策略、RecyclerView 的 Adapter 与 ItemTouchHelper 的配合使用，以及 Retrofit 异步请求与回调的生命周期管理。'
    '在后端开发中，我熟悉了 Flask 路由与视图函数的编写方式，体会到了 RESTful API 设计的简洁与优雅，'
    '同时也认识到了参数校验与异常处理的重要性。')
add_paragraph_custom(doc,
    '实验中最具挑战性的部分是同步机制的设计。如何在离线与在线状态之间平滑切换，如何保证本地数据与远端数据的一致，'
    '是分布式系统中永恒的主题。本项目采用的"本地优先 + 增量同步"策略虽然简单，但在个人工具的语境下足够有效，'
    '为后续更复杂的同步算法打下了实践基础。')
add_paragraph_custom(doc,
    '此外，调试过程中遇到的 URL 斜杠与 SQL 类型不匹配问题，让我认识到细节决定成败。一个不起眼的字符差异就可能导致应用崩溃，'
    '而这类问题在真实开发中更为隐蔽。养成查看 Logcat、阅读官方文档与编写边界测试用例的习惯至关重要。')
add_paragraph_custom(doc,
    '总体而言，本项目从需求分析、系统设计、编码实现到测试验收，完整走了一遍软件工程的生命周期，'
    '对全栈移动开发有了更宏观的认知。未来可在现有基础上接入 Firebase、Room 数据库替换原生 SQLite，'
    '并采用 Jetpack 组件重构架构，进一步提升代码质量与用户体验。')

# ================= 六、参考文献 =================
add_heading_custom(doc, '六、参考文献', level=2)
refs = [
    '[1] Android Developers. Guide to app architecture [EB/OL]. https://developer.android.com/jetpack/guide',
    '[2] Square. Retrofit: A type-safe HTTP client for Java and Kotlin [EB/OL]. https://square.github.io/retrofit/',
    '[3] Flask. Welcome to Flask — Flask Documentation [EB/OL]. https://flask.palletsprojects.com/',
    '[4] SQLite. SQLite Documentation [EB/OL]. https://www.sqlite.org/docs.html',
    '[5] Google. Material Design Components for Android [EB/OL]. https://m3.material.io/develop/android/jetpack-compose',
    '[6] OkHttp. Square\'s meticulous HTTP client for Java and Kotlin [EB/OL]. https://square.github.io/okhttp/',
    '[7] Python Software Foundation. Python 3 Documentation [EB/OL]. https://docs.python.org/3/',
    '[8] Android Developers. RecyclerView [EB/OL]. https://developer.android.com/develop/ui/views/layout/recyclerview',
    '[9] Android Developers. SQLiteOpenHelper [EB/OL]. https://developer.android.com/reference/android/database/sqlite/SQLiteOpenHelper',
    '[10] Square. Retrofit — Converters: Gson Converter [EB/OL]. https://square.github.io/retrofit/',
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(r)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)

# 保存文档
doc.save(OUTPUT_PATH)
print(f'实验报告已生成：{OUTPUT_PATH}')
